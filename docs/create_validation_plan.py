#!/usr/bin/env python3
"""
Generate the ENVISION Classifier Validation Plan (.docx) and
the Validation Review Template (.xlsx).
"""

import json
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

RESULTS_DIR = Path(__file__).resolve().parent.parent / "results"
DOCS_DIR = Path(__file__).resolve().parent


def add_heading_block(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * level)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    doc.add_paragraph()
    return table


def create_docx():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Title ──
    title = doc.add_heading("ENVISION Classifier Validation Plan", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("FAIR Data Innovations Hub — California Medical Innovations Institute (CalMI²)\n").bold = True
    meta.add_run(f"Prepared by: James O'Neill\n")
    meta.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}\n")
    meta.add_run("Status: DRAFT — for team review")
    doc.add_paragraph()

    # ── 1. Executive Summary ──
    add_heading_block(doc, "1. Executive Summary")
    add_body(doc,
        "The ENVISION eye imaging classifier identifies ophthalmic imaging datasets from "
        "Zenodo metadata using a SetFit few-shot model. To date, the classifier has been "
        "evaluated via a single-annotator spot check of 20 records. This plan defines the "
        "validation methodology required to achieve publication-grade evidence of classifier "
        "performance, including multi-reviewer annotation, inter-annotator agreement metrics, "
        "and per-class precision/recall/F1 scores."
    )

    # ── 2. Current State ──
    add_heading_block(doc, "2. Current State")

    add_heading_block(doc, "2.1 Classifier Output", level=2)
    add_table(doc,
        ["Class", "Count", "Description"],
        [
            ["EYE_IMAGING", "120", "Ophthalmic imaging datasets"],
            ["EYE_SOFTWARE", "66", "Code/tools for eye imaging"],
            ["OTHER_EYE_DATA", "3", "Ambiguous / borderline records"],
            ["NEGATIVE", "325", "Not eye-related"],
            ["Total", "514", "All classified Zenodo records"],
        ]
    )

    add_heading_block(doc, "2.2 Confidence Distribution (EYE_IMAGING)", level=2)
    add_table(doc,
        ["Tier", "Confidence", "Count", "% of EYE_IMAGING"],
        [
            ["High", "≥ 0.95", "117", "97.5%"],
            ["Medium", "0.80 – 0.95", "2", "1.7%"],
            ["Low", "< 0.80", "1", "0.8%"],
        ]
    )

    add_heading_block(doc, "2.3 Gap Analysis", level=2)
    add_bullet(doc, "Only 1 annotator (James) has reviewed any records")
    add_bullet(doc, "Only 20 of 514 records have been manually checked")
    add_bullet(doc, "No inter-annotator agreement (IAA) metrics exist")
    add_bullet(doc, "No confusion matrix or per-class F1 has been computed")
    add_bullet(doc, "No negative/edge-case records have been reviewed for false negatives")

    # ── 3. Validation Design ──
    add_heading_block(doc, "3. Validation Design")

    add_heading_block(doc, "3.1 Sampling Strategy", level=2)
    add_body(doc,
        "We will validate a stratified sample drawn from all four predicted classes to "
        "compute precision, recall, and F1 for each class. The sample must be large enough "
        "to produce meaningful confidence intervals."
    )
    add_table(doc,
        ["Predicted Class", "Population", "Sample Size", "Sampling Method"],
        [
            ["EYE_IMAGING", "120", "ALL 120", "Exhaustive — validate every record"],
            ["EYE_SOFTWARE", "66", "ALL 66", "Exhaustive — validate every record"],
            ["OTHER_EYE_DATA", "3", "ALL 3", "Exhaustive"],
            ["NEGATIVE", "325", "60", "Stratified random: 20 high-conf, 20 medium, 20 lowest-conf"],
            ["Total", "514", "249", "48% of full corpus"],
        ]
    )
    add_body(doc,
        "Rationale: All positive and software predictions are validated exhaustively (189 records). "
        "For negatives, we sample 60 to check for false negatives — focusing on lower-confidence "
        "negatives where misclassifications are most likely. This gives us a total of 249 records, "
        "a strong sample for publication."
    )

    add_heading_block(doc, "3.2 Reviewer Panel", level=2)
    add_body(doc, "A minimum of 3 independent reviewers is required for publication-grade validation:")
    add_table(doc,
        ["Reviewer", "Role", "Annotation Scope"],
        [
            ["James O'Neill", "Primary annotator + classifier developer", "All 249 records"],
            ["Sanjay Soundarajan", "Independent annotator", "All 249 records"],
            ["Dorian Portillo", "Independent annotator", "All 249 records"],
        ]
    )
    add_body(doc,
        "All 3 reviewers annotate all 249 records independently (full overlap). This enables "
        "pairwise and multi-rater agreement metrics. Reviewers must NOT see model predictions "
        "or confidence scores during annotation to avoid bias."
    )

    add_heading_block(doc, "3.3 Annotation Categories (Ground Truth Labels)", level=2)
    add_body(doc, "Each reviewer assigns exactly one label per record:")
    add_table(doc,
        ["Label", "Code", "Definition"],
        [
            ["EYE_IMAGING", "3", "Contains actual ophthalmic imaging data (OCT, fundus, OCTA, corneal, slit-lamp, etc.) with at least one downloadable data file"],
            ["EYE_SOFTWARE", "2", "Code, models, or tools for eye imaging analysis — no actual imaging data"],
            ["OTHER_EYE_DATA", "1", "Eye-related but not imaging data (GWAS, electrophysiology, reviews, animal non-imaging)"],
            ["NEGATIVE", "0", "Not related to eye/vision research at all"],
        ]
    )

    add_heading_block(doc, "3.4 Annotation Guidelines (per record)", level=2)
    add_body(doc, "For each Zenodo record, reviewers must:")
    add_bullet(doc, "Open the Zenodo URL and read the title and description")
    add_bullet(doc, "Check the \"Files\" section — inspect file names and types")
    add_bullet(doc, "For ambiguous cases, check ZIP contents (if listed) or download a sample")
    add_bullet(doc, "Assign one of the four labels above")
    add_bullet(doc, "Add free-text notes for any record that is difficult to classify")
    add_body(doc,
        "Estimated time per record: 1–3 minutes. Total per reviewer: ~4–12 hours."
    )

    # ── 4. Validation Interface ──
    add_heading_block(doc, "4. Validation Interface")

    add_heading_block(doc, "4.1 Spreadsheet-Based Review (Primary)", level=2)
    add_body(doc,
        "Each reviewer receives an Excel (.xlsx) workbook pre-populated with the 249 sampled "
        "records. The spreadsheet contains:"
    )
    add_bullet(doc, "Column A: Sequential number")
    add_bullet(doc, "Column B: Zenodo ID (hyperlinked to record URL)")
    add_bullet(doc, "Column C: Dataset title")
    add_bullet(doc, "Column D: Description (truncated to 300 chars)")
    add_bullet(doc, "Column E: File types detected")
    add_bullet(doc, "Column F: File count")
    add_bullet(doc, "Column G: Size (MB)")
    add_bullet(doc, "Column H: Reviewer label — data-validated dropdown (EYE_IMAGING / EYE_SOFTWARE / OTHER_EYE_DATA / NEGATIVE)")
    add_bullet(doc, "Column I: Confidence (1–5 scale) — how certain the reviewer is")
    add_bullet(doc, "Column J: Notes — free text for difficult cases")

    add_body(doc,
        "IMPORTANT: The spreadsheet must NOT include model predictions or confidence scores. "
        "These are withheld to prevent anchoring bias. Predictions are merged only after all "
        "reviewers submit their annotations."
    )

    add_heading_block(doc, "4.2 Template", level=2)
    add_body(doc,
        "An .xlsx template (envision_validation_template.xlsx) accompanies this document. "
        "Each reviewer receives their own copy. Completed spreadsheets are returned to James "
        "for agreement analysis."
    )

    # ── 5. Metrics & Analysis ──
    add_heading_block(doc, "5. Metrics & Analysis")

    add_heading_block(doc, "5.1 Inter-Annotator Agreement (IAA)", level=2)
    add_table(doc,
        ["Metric", "What It Measures", "Acceptable Threshold"],
        [
            ["Cohen's Kappa (pairwise)", "Agreement between each pair of annotators, corrected for chance", "κ ≥ 0.70 (substantial)"],
            ["Fleiss' Kappa (multi-rater)", "Agreement among all 3 annotators", "κ ≥ 0.70 (substantial)"],
            ["Krippendorff's Alpha", "Reliability across multiple coders (handles missing data)", "α ≥ 0.667 (tentative), α ≥ 0.80 (strong)"],
            ["Percent Agreement", "Raw agreement rate (no chance correction)", "Report but do not use as primary metric"],
        ]
    )

    add_heading_block(doc, "5.2 Establishing Ground Truth", level=2)
    add_body(doc, "After IAA is computed, ground truth labels are established by:")
    add_bullet(doc, "Majority vote: Label agreed upon by ≥ 2 of 3 reviewers")
    add_bullet(doc, "Disagreements (3-way split): Resolved by discussion and consensus, documented in adjudication log")

    add_heading_block(doc, "5.3 Classifier Performance (vs. Ground Truth)", level=2)
    add_table(doc,
        ["Metric", "Scope"],
        [
            ["Precision (per class)", "Of records the classifier labeled X, how many are truly X?"],
            ["Recall (per class)", "Of records that are truly X, how many did the classifier find?"],
            ["F1 Score (per class)", "Harmonic mean of precision and recall"],
            ["Macro F1", "Unweighted average F1 across all 4 classes"],
            ["Weighted F1", "F1 weighted by class support"],
            ["Confusion Matrix", "Full 4×4 matrix of predicted vs. actual labels"],
            ["Accuracy", "Overall correct / total (report but note class imbalance)"],
        ]
    )

    add_heading_block(doc, "5.4 Confidence Calibration", level=2)
    add_body(doc,
        "We will analyze whether classifier confidence correlates with actual correctness. "
        "Plot accuracy vs. confidence bins to produce a calibration curve. This is especially "
        "important because 97.5% of EYE_IMAGING predictions have confidence ≥ 0.95."
    )

    # ── 6. Timeline ──
    add_heading_block(doc, "6. Timeline")
    add_table(doc,
        ["Week", "Dates", "Task", "Owner"],
        [
            ["1", "Feb 17–21", "Finalize validation plan; generate & distribute .xlsx templates", "James"],
            ["2", "Feb 24–28", "Reviewers annotate 249 records (independent)", "Sanjay, Dorian, James"],
            ["3", "Mar 3–7", "Collect spreadsheets; compute IAA; adjudicate disagreements", "James"],
            ["4", "Mar 10–14", "Compute F1, confusion matrix, calibration; draft results section", "James"],
            ["5", "Mar 17–21", "Internal review of results; iterate if IAA < 0.70", "All"],
            ["6–8", "Mar 24 – Apr 11", "Write manuscript; prepare figures", "James + team"],
            ["9", "Apr 14–18", "Internal review and submission", "All"],
        ]
    )

    # ── 7. Target Journals ──
    add_heading_block(doc, "7. Target Journals")
    add_body(doc,
        "Ranked by fit for this work (automated dataset discovery + FAIR data + ophthalmology):"
    )
    add_table(doc,
        ["Priority", "Journal", "Publisher", "Why It Fits", "Impact / Notes"],
        [
            ["1", "Scientific Data", "Nature", "Purpose-built for dataset description papers; strong FAIR data alignment; our work discovers and catalogs datasets", "IF ~6.5; open access; fast review (~30 days)"],
            ["2", "JAMIA (J Am Med Inform Assoc)", "Oxford / AMIA", "Published DataMed (dataset discovery index); strong fit for biomedical informatics tools", "IF ~7.9; high visibility in medical informatics"],
            ["3", "Journal of Biomedical Informatics", "Elsevier / AMIA", "Methodology focus; NLP/ML classifiers for biomedical data; good fit for the classifier paper", "IF ~4.5; CiteScore 10.2"],
            ["4", "Data Science Journal", "CODATA", "FAIR data principles; dataset discovery tools; open science", "Open access; niche but well-regarded in data science"],
            ["5", "PLOS ONE", "PLOS", "Broad scope; accepts ML + medical data papers; open access", "IF ~3.7; fast turnaround; good fallback"],
            ["6", "MELBA (ML for Biomedical Imaging)", "Independent", "ML + medical imaging focus; open access; special issue on open data", "New journal; growing reputation"],
            ["7", "Artificial Intelligence in Vision & Ophthalmology (AIVO)", "SAIVO", "Directly targets AI + ophthalmology intersection", "New (2025); niche but perfect topic match"],
        ]
    )

    add_heading_block(doc, "7.1 Recommended Submission Strategy", level=2)
    add_body(doc,
        "Primary target: Scientific Data — frame the paper as a \"Data Descriptor\" for the "
        "ENVISION discovery corpus (the 120+ validated eye imaging datasets), with the classifier "
        "as the methodology that produced it. This journal explicitly welcomes papers describing "
        "curated dataset collections and discovery tools."
    )
    add_body(doc,
        "Secondary target: JAMIA or JBI — if we want to emphasize the classifier methodology "
        "and informatics contribution rather than the dataset itself."
    )

    # ── 8. Risks & Mitigations ──
    add_heading_block(doc, "8. Risks & Mitigations")
    add_table(doc,
        ["Risk", "Likelihood", "Impact", "Mitigation"],
        [
            ["Low IAA (κ < 0.70)", "Medium", "High — weakens publication", "Refine annotation guidelines; add training round with 10 practice records; re-annotate"],
            ["High false positive rate in EYE_IMAGING", "Low", "High", "Retrain classifier with new false positives as negative examples; re-classify"],
            ["Reviewer availability / delays", "Medium", "Medium", "Start ASAP; each reviewer needs ~6–12 hours over 1 week"],
            ["Zenodo records become unavailable", "Low", "Low", "Full metadata already cached locally; note any inaccessible records"],
            ["Class imbalance skews metrics", "High", "Medium", "Report both macro and weighted F1; discuss in limitations"],
        ]
    )

    # ── 9. Deliverables ──
    add_heading_block(doc, "9. Deliverables")
    add_bullet(doc, "3 completed annotation spreadsheets (.xlsx)")
    add_bullet(doc, "Inter-annotator agreement report (Cohen's κ, Fleiss' κ, Krippendorff's α)")
    add_bullet(doc, "Adjudication log for disagreements")
    add_bullet(doc, "Confusion matrix (4×4)")
    add_bullet(doc, "Per-class precision, recall, F1 table")
    add_bullet(doc, "Confidence calibration plot")
    add_bullet(doc, "Manuscript draft for target journal")

    out_path = DOCS_DIR / "ENVISION_Validation_Plan.docx"
    doc.save(str(out_path))
    print(f"Created: {out_path}")
    return out_path


def create_xlsx():
    """Create the validation template pre-populated with sampled records."""

    eye = json.loads((RESULTS_DIR / "zenodo_eye_imaging.json").read_text())
    sw = json.loads((RESULTS_DIR / "zenodo_software.json").read_text())
    all_results = json.loads((RESULTS_DIR / "zenodo_all_results.json").read_text())

    edge = [r for r in all_results if r["label"] == "OTHER_EYE_DATA"]
    neg = [r for r in all_results if r["label"] == "NEGATIVE"]

    neg_sorted = sorted(neg, key=lambda r: r["confidence"])
    neg_sample = []
    neg_sample.extend(neg_sorted[:20])       # 20 lowest confidence (most likely false negatives)
    mid = len(neg_sorted) // 2
    neg_sample.extend(neg_sorted[mid:mid+20]) # 20 medium confidence
    neg_sample.extend(neg_sorted[-20:])       # 20 highest confidence
    seen = set()
    neg_deduped = []
    for r in neg_sample:
        if r["zenodo_id"] not in seen:
            seen.add(r["zenodo_id"])
            neg_deduped.append(r)
    neg_sample = neg_deduped

    records = []
    for r in eye:
        records.append({**r, "_source": "EYE_IMAGING"})
    for r in sw:
        records.append({**r, "_source": "EYE_SOFTWARE"})
    for r in edge:
        records.append({**r, "_source": "OTHER_EYE_DATA"})
    for r in neg_sample:
        records.append({**r, "_source": "NEGATIVE (sampled)"})

    import random
    random.seed(42)
    random.shuffle(records)

    wb = openpyxl.Workbook()

    # ── Instructions sheet ──
    ws_inst = wb.active
    ws_inst.title = "Instructions"
    ws_inst.sheet_properties.tabColor = "4472C4"

    instructions = [
        ("ENVISION Classifier Validation", None),
        ("", None),
        ("Reviewer Name:", "[ENTER YOUR NAME]"),
        ("Date Started:", "[ENTER DATE]"),
        ("Date Completed:", "[ENTER DATE]"),
        ("", None),
        ("INSTRUCTIONS", None),
        ("1. Go to the 'Validation' tab", None),
        ("2. For each row, open the Zenodo URL (Column B, click the hyperlink)", None),
        ("3. Read the title, description, and inspect the Files section", None),
        ("4. In Column H, select your classification from the dropdown:", None),
        ("   - EYE_IMAGING: Contains actual ophthalmic imaging data (OCT, fundus, OCTA, corneal, etc.)", None),
        ("   - EYE_SOFTWARE: Code/models/tools for eye imaging — no actual data files", None),
        ("   - OTHER_EYE_DATA: Eye-related but not imaging (GWAS, genetics, electrophysiology, reviews)", None),
        ("   - NEGATIVE: Not related to eye/vision research", None),
        ("5. In Column I, rate your confidence (1=uncertain, 5=certain)", None),
        ("6. In Column J, add notes for difficult or ambiguous records", None),
        ("", None),
        ("IMPORTANT: Do NOT look at model predictions before annotating.", None),
        ("Your independent judgment is what makes this validation valid.", None),
        ("", None),
        (f"Total records to review: {len(records)}", None),
        ("Estimated time: 6–12 hours", None),
    ]
    for row_idx, (a, b) in enumerate(instructions, 1):
        ws_inst.cell(row=row_idx, column=1, value=a)
        if b:
            ws_inst.cell(row=row_idx, column=2, value=b)
    ws_inst.cell(row=1, column=1).font = Font(size=16, bold=True, color="1F4E79")
    ws_inst.cell(row=7, column=1).font = Font(size=13, bold=True, color="1F4E79")
    ws_inst.cell(row=19, column=1).font = Font(bold=True, color="C00000")
    ws_inst.cell(row=20, column=1).font = Font(bold=True, color="C00000")
    ws_inst.column_dimensions["A"].width = 80
    ws_inst.column_dimensions["B"].width = 30

    # ── Validation sheet ──
    ws = wb.create_sheet("Validation")
    ws.sheet_properties.tabColor = "70AD47"

    headers = [
        "#", "Zenodo ID", "Title", "Description", "File Types",
        "File Count", "Size (MB)", "Your Label", "Confidence (1-5)", "Notes"
    ]
    header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    for col_idx, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = thin_border

    col_widths = [5, 14, 55, 50, 18, 10, 10, 16, 14, 40]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    ws.auto_filter.ref = f"A1:J{len(records)+1}"
    ws.freeze_panes = "A2"

    alt_fill = PatternFill(start_color="F2F7FB", end_color="F2F7FB", fill_type="solid")

    for row_idx, r in enumerate(records, 2):
        zenodo_id = r.get("zenodo_id", "")
        url = r.get("url", f"https://zenodo.org/records/{zenodo_id}")
        title = r.get("title", "")[:120]
        desc = r.get("description", "")[:300]
        file_types = ", ".join(r.get("file_types", []))
        file_count = r.get("file_count", 0)
        size_mb = r.get("size_mb", 0)

        ws.cell(row=row_idx, column=1, value=row_idx - 1)
        id_cell = ws.cell(row=row_idx, column=2, value=zenodo_id)
        id_cell.hyperlink = url
        id_cell.font = Font(color="0563C1", underline="single")
        ws.cell(row=row_idx, column=3, value=title)
        ws.cell(row=row_idx, column=4, value=desc)
        ws.cell(row=row_idx, column=5, value=file_types)
        ws.cell(row=row_idx, column=6, value=file_count)
        ws.cell(row=row_idx, column=7, value=round(size_mb, 1))
        ws.cell(row=row_idx, column=8, value="")  # reviewer fills in
        ws.cell(row=row_idx, column=9, value="")  # reviewer fills in
        ws.cell(row=row_idx, column=10, value="")  # reviewer fills in

        if row_idx % 2 == 0:
            for col in range(1, 11):
                ws.cell(row=row_idx, column=col).fill = alt_fill

        for col in range(1, 11):
            ws.cell(row=row_idx, column=col).border = thin_border
            ws.cell(row=row_idx, column=col).alignment = Alignment(
                vertical="top", wrap_text=(col in [3, 4, 10])
            )

    from openpyxl.worksheet.datavalidation import DataValidation

    label_dv = DataValidation(
        type="list",
        formula1='"EYE_IMAGING,EYE_SOFTWARE,OTHER_EYE_DATA,NEGATIVE"',
        allow_blank=True,
    )
    label_dv.error = "Please select a valid label"
    label_dv.errorTitle = "Invalid Label"
    label_dv.prompt = "Select your classification"
    label_dv.promptTitle = "Label"
    ws.add_data_validation(label_dv)
    label_dv.add(f"H2:H{len(records)+1}")

    conf_dv = DataValidation(
        type="whole",
        operator="between",
        formula1="1",
        formula2="5",
        allow_blank=True,
    )
    conf_dv.error = "Enter a number 1–5"
    conf_dv.prompt = "1=uncertain, 5=very certain"
    ws.add_data_validation(conf_dv)
    conf_dv.add(f"I2:I{len(records)+1}")

    # ── Summary sheet (auto-calculated) ──
    ws_sum = wb.create_sheet("Summary")
    ws_sum.sheet_properties.tabColor = "ED7D31"

    ws_sum.cell(row=1, column=1, value="Validation Summary").font = Font(size=14, bold=True, color="1F4E79")
    ws_sum.cell(row=3, column=1, value="Label")
    ws_sum.cell(row=3, column=2, value="Count")
    for i, label in enumerate(["EYE_IMAGING", "EYE_SOFTWARE", "OTHER_EYE_DATA", "NEGATIVE", "(blank)"], 4):
        ws_sum.cell(row=i, column=1, value=label)
        if label == "(blank)":
            ws_sum.cell(row=i, column=2).value = f'=COUNTBLANK(Validation!H2:H{len(records)+1})'
        else:
            ws_sum.cell(row=i, column=2).value = f'=COUNTIF(Validation!H2:H{len(records)+1},"{label}")'

    ws_sum.cell(row=10, column=1, value="Total Annotated")
    ws_sum.cell(row=10, column=2).value = f'=COUNTA(Validation!H2:H{len(records)+1})'
    ws_sum.cell(row=11, column=1, value="Remaining")
    ws_sum.cell(row=11, column=2).value = f'=COUNTBLANK(Validation!H2:H{len(records)+1})'
    ws_sum.cell(row=12, column=1, value="% Complete")
    ws_sum.cell(row=12, column=2).value = f'=IF(COUNTA(Validation!H2:H{len(records)+1})=0,0,COUNTA(Validation!H2:H{len(records)+1})/{len(records)})'
    ws_sum.cell(row=12, column=2).number_format = "0.0%"

    for r in [3, 10, 11, 12]:
        ws_sum.cell(row=r, column=1).font = Font(bold=True)
    ws_sum.column_dimensions["A"].width = 20
    ws_sum.column_dimensions["B"].width = 15

    out_path = DOCS_DIR / "envision_validation_template.xlsx"
    wb.save(str(out_path))
    print(f"Created: {out_path}  ({len(records)} records)")
    return out_path


if __name__ == "__main__":
    create_docx()
    create_xlsx()
    print("\nDone. Files saved to docs/")
