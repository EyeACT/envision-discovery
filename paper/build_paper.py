#!/usr/bin/env python3
"""Build a publication-quality .docx for the Envision Discovery paper."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ── Style definitions ───────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

# Heading 1 — section headings
h1 = doc.styles['Heading 1']
h1.font.name = 'Times New Roman'
h1.font.size = Pt(14)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0, 0, 0)
h1.paragraph_format.space_before = Pt(24)
h1.paragraph_format.space_after = Pt(6)
h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Heading 2 — subsections
h2 = doc.styles['Heading 2']
h2.font.name = 'Times New Roman'
h2.font.size = Pt(12)
h2.font.bold = True
h2.font.italic = False
h2.font.color.rgb = RGBColor(0, 0, 0)
h2.paragraph_format.space_before = Pt(18)
h2.paragraph_format.space_after = Pt(4)
h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

# Heading 3 — sub-subsections
h3 = doc.styles['Heading 3']
h3.font.name = 'Times New Roman'
h3.font.size = Pt(12)
h3.font.bold = False
h3.font.italic = True
h3.font.color.rgb = RGBColor(0, 0, 0)
h3.paragraph_format.space_before = Pt(12)
h3.paragraph_format.space_after = Pt(4)
h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE


def add_text(text, bold=False, italic=False, size=None, alignment=None,
             space_after=None, first_line_indent=None, style_name='Normal'):
    """Add a paragraph with optional formatting."""
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    if size:
        run.font.size = Pt(size)
    if alignment is not None:
        p.paragraph_format.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Inches(first_line_indent)
    return p


def add_rich_paragraph(parts, alignment=None, space_after=None,
                       space_before=None, first_line_indent=None):
    """Add paragraph with mixed bold/italic/normal runs.
    parts: list of (text, bold, italic) tuples
    """
    p = doc.add_paragraph()
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    if alignment is not None:
        p.paragraph_format.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Inches(first_line_indent)
    return p


def add_body(text, first_line_indent=0.5):
    """Standard body paragraph with first-line indent."""
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Inches(first_line_indent)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p


def add_table(headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        # Shade header
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    return table


def add_bullet(text, bold_prefix=None):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    else:
        p.runs[0].font.name = 'Times New Roman' if p.runs else None
        if not bold_prefix:
            p.clear()
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    return p


def add_numbered(text, number, bold_prefix=None):
    """Add a numbered item."""
    p = doc.add_paragraph()
    prefix_text = f"{number}. "
    if bold_prefix:
        run = p.add_run(prefix_text + bold_prefix)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    else:
        run = p.add_run(prefix_text + text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    return p


# ════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ════════════════════════════════════════════════════════════════════

# Add some vertical space
add_text('', space_after=36)

add_text(
    'Envision Discovery: Automated Identification of Eye Imaging Datasets '
    'Across Scientific Repositories',
    bold=True, size=16,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_after=18
)

add_text('', space_after=6)

add_rich_paragraph([
    ('James O\'Neill', False, False),
    ('1', False, False),
    (', [Validator Names]', False, False),
    ('2', False, False),
    (', Bhavesh Patel', False, False),
    ('1', False, False),
], alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

# Superscript the affiliation numbers manually
for p in doc.paragraphs:
    for run in p.runs:
        pass  # We'll handle superscripts below

add_text(
    '\u00b9 FAIR Data Innovations Hub, California Medical Innovations Institute (CalMI\u00b2), '
    'San Diego, CA, USA',
    size=10, italic=True,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_after=2
)

add_text(
    '\u00b2 [Validator Affiliations]',
    size=10, italic=True,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_after=12
)

add_text('', space_after=4)

add_rich_paragraph([
    ('Correspondence: ', True, False),
    ('joneill@calmi2.org, bpatel@calmi2.org', False, False),
], alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

# ════════════════════════════════════════════════════════════════════
# ABSTRACT
# ════════════════════════════════════════════════════════════════════

doc.add_heading('Abstract', level=1)

abstract = (
    'Eye imaging datasets\u2014including optical coherence tomography (OCT), fundus photography, '
    'and OCT angiography (OCTA)\u2014are essential resources for developing artificial intelligence '
    '(AI) tools in ophthalmology. However, these datasets are scattered across generalist '
    'repositories with no centralized catalog, making discovery and reuse prohibitively difficult. '
    'Here we present Envision Discovery, a machine learning pipeline that automatically identifies '
    'eye imaging datasets from scientific data repositories. The system uses a SetFit few-shot '
    'classifier built on a large language embedding model (GTE-large, 1024-dimensional) trained on '
    '452 curated examples to distinguish four classes: genuine eye imaging datasets, eye-related '
    'software, edge cases, and unrelated records. Applied to 30,439 metadata records harvested from '
    'Zenodo, the pipeline identified 120 candidate eye imaging datasets from 514 records containing '
    'data files, with 97.5% of positive predictions at high confidence (\u22650.95). Expert validation '
    'by [N] ophthalmology specialists across [N] independent batches confirmed a precision of [XX]% '
    'for high-confidence predictions. All identified datasets are registered on the Envision Portal '
    '(https://envisionportal.org), providing researchers with a single entry point for discovering '
    'publicly available ophthalmic imaging data. The classifier, pipeline code, and trained model '
    'are openly available at https://github.com/EyeACT/envision-discovery.'
)
add_body(abstract, first_line_indent=0)

add_text('', space_after=4)
add_rich_paragraph([
    ('Keywords: ', True, False),
    ('eye imaging, dataset discovery, machine learning, FAIR data, ophthalmology, OCT, '
     'fundus photography, SetFit, few-shot learning', False, True),
], space_after=6)


# ════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ════════════════════════════════════════════════════════════════════

doc.add_heading('1. Introduction', level=1)

add_body(
    'Eye imaging modalities such as OCT, OCTA, fundus photography, and fluorescence lifetime '
    'imaging ophthalmoscopy (FLIO) provide detailed structural and functional views of ocular '
    'tissues and have become indispensable for ophthalmic research and clinical innovation. These '
    'datasets have been particularly critical for developing AI-based diagnostic tools. RETFound, '
    'a self-supervised foundation model trained on 1.6 million unlabeled retinal images, '
    'demonstrated state-of-the-art performance across both ophthalmic and systemic disease '
    'detection tasks (Zhou et al., 2023). OphGLM combined fundus photography with natural language '
    'capabilities for diagnostic support (OphGLM, 2024), while Ophtha-LLaMA2 fine-tuned a large '
    'language model on multimodal ophthalmic data for efficient clinical deployment (Zhao et al., '
    '2023). These advances depend critically on the availability and accessibility of large-scale '
    'imaging datasets.'
)

add_body(
    'However, sharing and discovering eye imaging data remains a significant challenge. Unlike '
    'neuroimaging, where community standards such as BIDS (Gorgolewski et al., 2016) and '
    'centralized platforms like OpenNeuro (Markiewicz et al., 2021) have transformed data sharing '
    'practices, ophthalmology lacks equivalent infrastructure. Eye imaging datasets are fragmented '
    'across generalist repositories\u2014Zenodo, Figshare, Dryad, and institutional archives\u2014that do '
    'not support ophthalmic-specific metadata or enforce standardized data formatting. This '
    'fragmentation makes datasets difficult to find and, when found, often difficult to reuse '
    '(Gim et al., 2025).'
)

add_body(
    'We experienced this challenge directly when searching for open-access OCT data related to '
    'age-related macular degeneration (AMD). Despite extensive multi-platform searches using '
    'diverse query strategies, finding relevant datasets was time-consuming and often required '
    'reading through published manuscripts on PubMed to locate associated data deposits (Gim et '
    'al., 2025). The absence of a centralized, searchable catalog means that researchers routinely '
    'duplicate collection efforts or remain unaware of existing data that could accelerate their '
    'work.'
)

add_body(
    'To address this gap, we developed Envision Discovery, an automated pipeline that harvests '
    'metadata from scientific data repositories, classifies records using a few-shot machine '
    'learning model, and identifies those containing eye imaging data. The identified datasets are '
    'then registered on the Envision Portal (https://envisionportal.org), a platform developed as '
    'part of the Eye Aging, Cognition, and Imaging (EyeACT) study to share and discover FAIR '
    '(Findable, Accessible, Interoperable, Reusable) and AI-ready eye imaging datasets. In this '
    'paper, we describe the design and implementation of the Envision Discovery pipeline, present '
    'results from its application to the Zenodo repository, and report on expert validation of '
    'classifier predictions by ophthalmology domain specialists.'
)


# ════════════════════════════════════════════════════════════════════
# 2. METHODS
# ════════════════════════════════════════════════════════════════════

doc.add_heading('2. Methods', level=1)

doc.add_heading('2.1 Overview', level=2)

add_body(
    'The Envision Discovery pipeline consists of four stages: (1) metadata harvesting from '
    'scientific repositories, (2) intelligent filtering and enrichment, (3) automated '
    'classification using a few-shot learning model, and (4) expert validation. Figure 1 provides '
    'an overview of the pipeline architecture.'
)

add_text(
    '[Figure 1. Pipeline overview: metadata harvesting \u2192 filtering & enrichment \u2192 '
    'classification \u2192 expert validation \u2192 Envision Portal registration.]',
    italic=True,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    space_after=6
)

doc.add_heading('2.2 Metadata Harvesting', level=2)

add_body(
    'We developed a scraper targeting the Zenodo repository (https://zenodo.org), a '
    'general-purpose open repository operated by CERN that hosts over 3 million research outputs. '
    'The scraper queries the Zenodo REST API using 249 unique search terms spanning:'
)

add_bullet('OCT, OCTA, SD-OCT, SS-OCT, AS-OCT, fundus photography, fluorescein angiography, '
           'indocyanine green angiography (ICGA), slit-lamp biomicroscopy, confocal microscopy, '
           'adaptive optics, photoacoustic imaging',
           bold_prefix='Imaging modalities: ')

add_bullet('retina, macula, fovea, optic nerve/disc, choroid, cornea, lens, iris, anterior chamber',
           bold_prefix='Anatomical structures: ')

add_bullet('retinal nerve fiber layer (RNFL), ganglion cell layer (GCL), retinal pigment '
           'epithelium (RPE), ellipsoid zone',
           bold_prefix='Retinal layers: ')

add_bullet('diabetic retinopathy, glaucoma, AMD, diabetic macular edema (DME), retinal '
           'detachment, retinitis pigmentosa, Stargardt disease, keratoconus, retinopathy of '
           'prematurity (ROP)',
           bold_prefix='Diseases: ')

add_bullet('Heidelberg Spectralis, Zeiss Cirrus, Topcon Triton/DRI OCT, Optovue RTVue, Optos',
           bold_prefix='Equipment and vendors: ')

add_bullet('DRIVE, STARE, CHASE_DB1, MESSIDOR, IDRiD, APTOS, REFUGE, EyePACS, AIROGS',
           bold_prefix='Benchmark datasets: ')

add_bullet('visual acuity, LogMAR, RNFL thickness, central macular thickness (CMT), cup-to-disc '
           'ratio, vessel density, foveal avascular zone (FAZ) area',
           bold_prefix='Clinical measurements: ')

add_body(
    'Search results are filtered to include only records with resource_type=dataset, excluding '
    'publications, software, and miscellaneous uploads.'
)

doc.add_heading('2.3 Filtering and Metadata Enrichment', level=2)

add_body('Retrieved records undergo several enrichment steps:', first_line_indent=0)

add_rich_paragraph([
    ('File-type filtering. ', True, False),
    ('Records are retained if they contain recognized data files (.dcm, .nii, .nii.gz, .mat, '
     '.h5, .hdf5, .npy, .npz, .jpg, .jpeg, .png, .tif, .tiff, .bmp) or archives (.zip, '
     '.tar.gz, .rar, .7z). Records containing only genomics files (.fasta, .h5ad, .vcf, .bam, '
     '.fastq) are excluded to reduce false positives from genomics studies that reference '
     'ophthalmic phenotypes.', False, False),
], first_line_indent=0.5)

add_rich_paragraph([
    ('ZIP content inspection. ', True, False),
    ('For archived records, we implemented a non-destructive inspection technique using HTTP '
     'Range requests. By downloading only the ZIP central directory (the last ~64 KB of the '
     'archive), we extract the complete file manifest without downloading the full archive. This '
     'catalogued 31,958 files across all inspected archives and enabled accurate file-type '
     'profiling at minimal bandwidth cost.', False, False),
], first_line_indent=0.5)

add_rich_paragraph([
    ('Link extraction. ', True, False),
    ('External dataset links are extracted from Zenodo\u2019s related_identifiers metadata field '
     'and from description text, capturing references to platforms including GitHub, Kaggle, '
     'HuggingFace, Google Drive, OSF, Dryad, Figshare, and others. Links are categorized as '
     'data platform, archive download, direct file, or potential download.', False, False),
], first_line_indent=0.5)

add_rich_paragraph([
    ('Metadata normalization. ', True, False),
    ('For each record, we extract and normalize: title, description (HTML stripped), keywords, '
     'file types, file counts, total size, and counts of imaging, medical, archive, and genomics '
     'files.', False, False),
], first_line_indent=0.5)


doc.add_heading('2.4 Classification Model', level=2)

doc.add_heading('2.4.1 Architecture', level=3)

add_body(
    'We use SetFit (Tunstall et al., 2022), a few-shot text classification framework that '
    'combines contrastive learning on sentence embeddings with a lightweight classification head. '
    'SetFit is well-suited to our task because it achieves strong performance with limited labeled '
    'data\u2014critical given the specialized nature of ophthalmic dataset descriptions.'
)

add_body(
    'The backbone model is Alibaba-NLP/gte-large-en-v1.5, a general text embedding model '
    'producing 1024-dimensional representations with an 8,192-token context window. It uses 24 '
    'transformer layers with 16 attention heads and Rotary Position Embeddings (RoPE) with NTK '
    'scaling for long-context handling. The classification head is a logistic regression model '
    'trained on the contrastively learned embeddings.'
)

doc.add_heading('2.4.2 Classification Schema', level=3)

add_body(
    'Records are classified into four mutually exclusive categories (Table 1).',
    first_line_indent=0
)

add_text('', space_after=4)

add_table(
    headers=['Class', 'Label', 'Description'],
    rows=[
        ['3', 'EYE_IMAGING', 'Contains actual ophthalmic imaging data (e.g., OCT scans, fundus photographs, OCTA images, corneal imaging, slit-lamp photographs)'],
        ['2', 'EYE_SOFTWARE', 'Code, tools, or pre-trained models related to eye imaging analysis, but no actual image data'],
        ['1', 'EDGE_CASE', 'Related to eye or vision research but not imaging datasets (e.g., genetics/GWAS, electrophysiology, eye tracking, animal models)'],
        ['0', 'NEGATIVE', 'Unrelated to eye or vision research'],
    ],
    col_widths=[0.6, 1.3, 4.6]
)

add_rich_paragraph([
    ('Table 1. ', True, False),
    ('Four-class classification schema for the Envision Discovery classifier.', False, True),
], alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=12)

add_body(
    'This four-class schema was designed to address specific false-positive patterns observed '
    'during development. For example, cardiovascular OCT (intravascular OCT of coronary arteries), '
    'industrial OCT (semiconductor inspection), and dental OCT are classified as NEGATIVE despite '
    'sharing terminology with ophthalmic OCT. Software repositories containing eye imaging analysis '
    'tools but no image data are common on Zenodo and are captured by the EYE_SOFTWARE class. GWAS '
    'studies referencing retinal phenotypes and electrophysiology studies (ERG, VEP) are captured '
    'as EDGE_CASE.'
)

doc.add_heading('2.4.3 Training Data', level=3)

add_body(
    'The training set consists of 452 manually curated examples (Table 2).',
    first_line_indent=0
)

add_text('', space_after=4)

add_table(
    headers=['Class', 'Count', 'Proportion', 'Representative Examples'],
    rows=[
        ['EYE_IMAGING', '99', '21.9%', 'IDRiD, REFUGE, RFMiD, OLIVES, Rotterdam EyePACS, retinal vessel segmentation datasets'],
        ['EYE_SOFTWARE', '30', '6.6%', 'GitHub repositories, segmentation model weights, Python/MATLAB packages, ImageJ plugins'],
        ['EDGE_CASE', '90', '19.9%', 'DR detection review papers, glaucoma ML literature, GWAS meta-analyses, zebrafish/Drosophila eye development, eye tracking'],
        ['NEGATIVE', '233', '51.5%', 'Climate data, COVID genomics, face recognition (LFW), cardiac imaging, brain MRI, MNIST, robotics, taxonomy papers'],
    ],
    col_widths=[1.1, 0.6, 0.8, 4.0]
)

add_rich_paragraph([
    ('Table 2. ', True, False),
    ('Distribution of training examples across the four classification categories.', False, True),
], alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=12)

add_body(
    'Training examples were curated to cover known confounding patterns, with particular attention '
    'to negative examples that share superficial similarity with eye imaging (e.g., intravascular '
    'OCT, hand-eye calibration in robotics, taxonomy papers with figure references). The model was '
    'trained for 2 epochs with a batch size of 16.'
)

doc.add_heading('2.4.4 Input Representation', level=3)

add_body(
    'For each record, the classifier receives a concatenation of the title, description (with '
    'HTML tags removed), and keywords as a single text input. This combined representation '
    'provides sufficient context for the embedding model to capture both the domain specificity '
    '(ophthalmology vs. other fields) and the data-type specificity (imaging data vs. software '
    'vs. publications).'
)

doc.add_heading('2.5 Expert Validation Protocol', level=2)

add_body(
    'To rigorously evaluate classifier performance, we designed a multi-validator expert review '
    'protocol. Ophthalmology domain specialists independently review classifier predictions '
    'through a dedicated web interface.'
)

doc.add_heading('2.5.1 Validation Interface', level=3)

add_body(
    'Each validator is presented with dataset records including the dataset title, a brief '
    'description or abstract, the file types contained in the record, and a link to the original '
    'Zenodo record. For each record, validators assign one of the four classification labels (Eye '
    'Imaging, Eye Software, Edge Case, Not Related) and a confidence score from 0 to 5 indicating '
    'their certainty in the assessment.'
)

doc.add_heading('2.5.2 Validation Batches', level=3)

add_body(
    'Records are distributed to validators in batches of 50. Each batch is designed to include a '
    'mix of high-confidence and lower-confidence classifier predictions to evaluate performance '
    'across the confidence spectrum. Multiple validators review overlapping subsets to enable '
    'inter-rater agreement analysis. Each batch is estimated to require 1\u20133 hours of reviewer '
    'time.'
)

doc.add_heading('2.5.3 Evaluation Metrics', level=3)

add_body('Classifier performance is evaluated using:', first_line_indent=0)

add_bullet(
    'proportion of EYE_IMAGING predictions that are confirmed by experts',
    bold_prefix='Precision (positive predictive value): '
)
add_bullet(
    'proportion of true eye imaging datasets correctly identified by the classifier, '
    'assessed via a held-out sample of manually curated records',
    bold_prefix='Recall (sensitivity): '
)
add_bullet(
    'Cohen\u2019s kappa or Fleiss\u2019 kappa for overlapping validation batches',
    bold_prefix='Inter-rater agreement: '
)
add_bullet(
    'correlation between classifier confidence scores and expert agreement rates',
    bold_prefix='Confidence calibration: '
)


# ════════════════════════════════════════════════════════════════════
# 3. RESULTS
# ════════════════════════════════════════════════════════════════════

doc.add_heading('3. Results', level=1)

doc.add_heading('3.1 Metadata Harvesting', level=2)

add_body(
    'The scraper retrieved metadata for 30,439 records from Zenodo matching at least one of the '
    '249 search terms. After file-type filtering, 514 records (1.7%) contained recognized data '
    'files or archives and were retained for classification. The low retention rate reflects the '
    'predominance of publications, software, and non-data records in Zenodo search results, even '
    'when filtering for resource_type=dataset.'
)

doc.add_heading('3.2 Classification Results', level=2)

add_body(
    'Of the 514 filtered records, the classifier produced the following distribution (Table 3).',
    first_line_indent=0
)

add_text('', space_after=4)

add_table(
    headers=['Class', 'Count', 'Proportion'],
    rows=[
        ['NEGATIVE', '325', '63.2%'],
        ['EYE_IMAGING', '120', '23.3%'],
        ['EYE_SOFTWARE', '66', '12.8%'],
        ['EDGE_CASE', '3', '0.6%'],
    ],
    col_widths=[2.0, 1.5, 1.5]
)

add_rich_paragraph([
    ('Table 3. ', True, False),
    ('Classification distribution across 514 filtered Zenodo records.', False, True),
], alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=12)

add_body(
    'The 120 EYE_IMAGING predictions showed notably high confidence: 117 (97.5%) had confidence '
    'scores \u22650.95, 2 (1.7%) fell in the 0.80\u20130.95 range, and only 1 (0.8%) was below 0.80. '
    'This bimodal confidence distribution\u2014with most predictions either very high or very low '
    'confidence\u2014suggests the classifier learned well-separated class boundaries.'
)

add_body(
    'The identified eye imaging datasets span a total volume of approximately 489.4 GB, with '
    'individual datasets ranging from a few megabytes (segmentation benchmarks) to 37.7 GB (Human '
    'Developing Retina Atlas). Represented modalities include OCT, OCTA, fundus photography, '
    'corneal imaging, and slit-lamp photography.'
)

doc.add_heading('3.3 Preliminary Spot-Check Validation', level=2)

add_body(
    'Prior to formal expert validation, a preliminary spot-check of the top 20 high-confidence '
    'predictions was conducted. Of 17 records evaluated, 14 (82%) were confirmed as valid eye '
    'imaging datasets, 3 (18%) required further review due to ambiguity (e.g., code-with-data '
    'hybrids, GWAS studies with retinal phenotypes), and none were false positives.'
)

add_rich_paragraph([
    ('[Table 4. Spot-check validation of the top 20 high-confidence EYE_IMAGING predictions, '
     'including Zenodo ID, title, dataset size, confidence score, and validation status.]',
     False, True),
], alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=8)

add_body('Notable discoveries include:', first_line_indent=0)

add_bullet(
    '3 related records totaling over 65 GB of multi-modal retinal imaging data',
    bold_prefix='Human Developing Retina Atlas: '
)
add_bullet(
    'Over 70 GB across multiple studies using Spectralis, Zeiss, and Topcon instruments',
    bold_prefix='Corneal OCT collections: '
)
add_bullet(
    'RTN4IP1 optic atrophy (5.2 GB), polypoidal choroidal vasculopathy/choroidal '
    'neovascularization imaging',
    bold_prefix='Clinical imaging studies: '
)
add_bullet(
    'nnUNet optic disc segmentation, OCTSEG, retinal vessel datasets',
    bold_prefix='Segmentation benchmarks: '
)

doc.add_heading('3.4 Expert Validation Results', level=2)

add_rich_paragraph([
    ('[This section will be populated following completion of expert validation batches.]',
     False, True),
], space_after=12)

add_body(
    '[N] ophthalmology specialists validated [N] records across [N] batches of 50. Overall results:'
)

add_bullet('[XX]%', bold_prefix='Precision for EYE_IMAGING (high confidence, \u22650.95): ')
add_bullet('[XX]%', bold_prefix='Precision for EYE_IMAGING (all confidence levels): ')
add_bullet('[X.XX]', bold_prefix='Inter-rater agreement (Cohen\u2019s kappa): ')
add_bullet('[XX]%', bold_prefix='False positive rate: ')
add_bullet('[to be determined]', bold_prefix='Common misclassification patterns: ')

add_rich_paragraph([
    ('[Table 5. Expert validation confusion matrix across all validated records.]',
     False, True),
], alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)

add_rich_paragraph([
    ('[Figure 2. Classifier confidence vs. expert agreement rate, demonstrating calibration '
     'of confidence scores.]',
     False, True),
], alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)


# ════════════════════════════════════════════════════════════════════
# 4. DISCUSSION
# ════════════════════════════════════════════════════════════════════

doc.add_heading('4. Discussion', level=1)

doc.add_heading('4.1 Automated Dataset Discovery as Infrastructure', level=2)

add_body(
    'Envision Discovery demonstrates the feasibility of using few-shot machine learning to '
    'automate the identification of domain-specific datasets from generalist repositories. The '
    '452-example training set\u2014small by typical machine learning standards\u2014proved sufficient to '
    'learn discriminative features for ophthalmic imaging datasets when combined with the strong '
    'pre-trained representations of GTE-large and the contrastive learning framework of SetFit. '
    'This few-shot approach is particularly advantageous for specialized scientific domains where '
    'large labeled datasets are impractical to construct.'
)

add_body(
    'The four-class schema proved essential for practical deployment. A binary (eye imaging vs. '
    'not) classifier would conflate software tools, edge cases, and truly unrelated records into '
    'a single negative class, reducing interpretability and making it difficult to identify '
    'improvement targets. The EYE_SOFTWARE class, in particular, captures a substantial category '
    '(12.8% of filtered records) that is relevant to the eye imaging community but should not be '
    'presented as imaging data.'
)

doc.add_heading('4.2 Confidence-Based Triaging', level=2)

add_body(
    'The extreme confidence distribution\u201497.5% of positive predictions above 0.95\u2014suggests the '
    'classifier effectively separates clear positives from ambiguous cases. This enables a '
    'confidence-based triaging strategy for the Envision Portal: high-confidence predictions can '
    'be surfaced immediately with minimal review overhead, while lower-confidence predictions are '
    'prioritized for expert validation. Such a strategy can scale dataset discovery to much larger '
    'repository collections without proportionally increasing validation burden.'
)

doc.add_heading('4.3 Challenges and False Positive Patterns', level=2)

add_body(
    'Several categories of records posed classification challenges:',
    first_line_indent=0
)

add_bullet(
    'Cardiovascular, dermatological, and industrial OCT share imaging terminology but are '
    'unrelated to ophthalmology. The training set includes explicit negative examples for these '
    'domains.',
    bold_prefix='Non-ophthalmic OCT. '
)
add_bullet(
    'Some software repositories include small sample datasets alongside code. These fall in a '
    'gray area between EYE_SOFTWARE and EYE_IMAGING.',
    bold_prefix='Code-with-data hybrids. '
)
add_bullet(
    'GWAS and genetic studies that reference retinal measurements (e.g., RNFL thickness as a '
    'phenotype) are clinically adjacent but do not contain imaging data.',
    bold_prefix='Genetics with retinal phenotypes. '
)
add_bullet(
    'Some Zenodo records aggregate heterogeneous content where eye imaging is a minor component.',
    bold_prefix='Multi-domain repositories. '
)

doc.add_heading('4.4 Comparison with Manual Discovery', level=2)

add_body(
    'The contrast between automated and manual dataset discovery is stark. Gim et al. (2025) '
    'documented the extensive effort required to manually locate AMD-related OCT datasets: '
    'multi-platform searches, diverse query strategies, manuscript reading on PubMed, and still '
    'only a small number of datasets meeting eligibility criteria. Envision Discovery screened '
    '30,439 records and identified 120 eye imaging datasets in approximately 30 minutes of '
    'compute time\u2014a task that would take weeks or months of manual effort and would inevitably '
    'miss datasets described with non-standard terminology.'
)

doc.add_heading('4.5 Limitations', level=2)

add_numbered(
    ' The current pipeline targets only Zenodo. Extension to Figshare, Dryad, OSF, and '
    'institutional repositories is planned and will substantially increase coverage.',
    1, bold_prefix='Single repository.'
)
add_numbered(
    ' Search terms and training examples are in English. Datasets described in other languages '
    'may be missed.',
    2, bold_prefix='Language bias.'
)
add_numbered(
    ' The classifier operates on metadata (titles, descriptions, keywords) without examining '
    'actual data files. Records with uninformative metadata may be misclassified.',
    3, bold_prefix='Metadata-only classification.'
)
add_numbered(
    ' Content behind access controls cannot be independently verified without requesting access.',
    4, bold_prefix='Restricted-access records.'
)
add_numbered(
    ' New deposits, updated metadata, and removed records require periodic re-scraping to '
    'maintain currency.',
    5, bold_prefix='Evolving repositories.'
)
add_numbered(
    ' While 452 examples proved sufficient for the current task, performance on rare or novel '
    'dataset types may be limited.',
    6, bold_prefix='Training set size.'
)

doc.add_heading('4.6 Future Directions', level=2)

add_body('Immediate next steps include:', first_line_indent=0)

add_bullet(
    'Adapting the scraper for Figshare, Dryad, OSF, and DataCite Commons',
    bold_prefix='Multi-repository expansion: '
)
add_bullet(
    'Using expert validation results to iteratively refine the training set and retrain the '
    'classifier',
    bold_prefix='Active learning: '
)
add_bullet(
    'Developing automated metrics for metadata completeness, documentation quality, and '
    'AI-readiness',
    bold_prefix='Dataset quality scoring: '
)
add_bullet(
    'Implementing automated monthly re-scraping with incremental classification updates',
    bold_prefix='Continuous monitoring: '
)
add_bullet(
    'Expanding search terms and training examples to support non-English datasets',
    bold_prefix='Multilingual support: '
)


# ════════════════════════════════════════════════════════════════════
# 5. CONCLUSION
# ════════════════════════════════════════════════════════════════════

doc.add_heading('5. Conclusion', level=1)

add_body(
    'We present Envision Discovery, an automated pipeline for identifying eye imaging datasets '
    'across scientific repositories. By combining targeted metadata harvesting, intelligent '
    'filtering with non-destructive archive inspection, and few-shot classification using SetFit, '
    'the system identified 120 eye imaging datasets from over 30,000 Zenodo records with high '
    'confidence and [XX]% expert-validated precision. The discovered datasets, spanning '
    'approximately 489 GB across OCT, OCTA, fundus photography, and other ophthalmic modalities, '
    'are now catalogued on the Envision Portal, providing the research community with a '
    'centralized resource for finding publicly available eye imaging data. As the pipeline extends '
    'to additional repositories and incorporates feedback from expert validation, Envision '
    'Discovery will serve as a continuously expanding, community-validated catalog to accelerate '
    'AI development and clinical research in ophthalmology.'
)


# ════════════════════════════════════════════════════════════════════
# DATA AND CODE AVAILABILITY
# ════════════════════════════════════════════════════════════════════

doc.add_heading('Data and Code Availability', level=1)

add_body(
    'The pipeline code and trained model are available at '
    'https://github.com/EyeACT/envision-discovery. The model is published on HuggingFace at '
    'https://huggingface.co/fairdataihub/envision-eye-imaging-classifier. The Envision Portal is '
    'accessible at https://envisionportal.org. All classified results are available in the '
    'repository under results/.',
    first_line_indent=0
)


# ════════════════════════════════════════════════════════════════════
# ACKNOWLEDGMENTS
# ════════════════════════════════════════════════════════════════════

doc.add_heading('Acknowledgments', level=1)

add_body(
    'This work was supported by [NIH grant number]. The Envision Portal is developed as part of '
    'the Eye Aging, Cognition, and Imaging (EyeACT) study (https://eyeactstudy.org). We thank '
    '[validator names] for their expert validation of classifier predictions. [Additional '
    'acknowledgments.]',
    first_line_indent=0
)


# ════════════════════════════════════════════════════════════════════
# REFERENCES
# ════════════════════════════════════════════════════════════════════

doc.add_heading('References', level=1)

refs = [
    'Gim, N., et al. (2025). Publicly available imaging datasets for age-related macular '
    'degeneration: Evaluation according to the Findable, Accessible, Interoperable, Reusable '
    '(FAIR) principles. Experimental Eye Research, 255, 110342.',

    'Gorgolewski, K. J., et al. (2016). The brain imaging data structure, a format for '
    'organizing and describing outputs of neuroimaging experiments. Scientific Data, 3, 160044.',

    'Markiewicz, C. J., et al. (2021). The OpenNeuro resource for sharing of neuroscience data. '
    'eLife, 10, e71774.',

    'OphGLM. (2024). OphGLM: An ophthalmology large language-and-vision assistant. Artificial '
    'Intelligence in Medicine, 157, 103001.',

    'Patoni, S. I. P., et al. (2023). Artificial intelligence in ophthalmology. Romanian Journal '
    'of Ophthalmology, 67, 207.',

    'Tan, Y. Y., et al. (2024). Prognostic potentials of AI in ophthalmology: systemic disease '
    'forecasting via retinal imaging. Eye and Vision, 11, 1\u201318.',

    'Tunstall, L., et al. (2022). Efficient few-shot learning without prompts. '
    'arXiv:2209.11055.',

    'Wang, Z., et al. (2022). Artificial Intelligence and Deep Learning in Ophthalmology. In '
    'Artificial Intelligence in Medicine (pp. 1519\u20131552).',

    'Zhao, H., et al. (2023). Ophtha-LLaMA2: A Large Language Model for Ophthalmology. arXiv '
    'preprint.',

    'Zhou, Y., et al. (2023). A foundation model for generalizable disease detection from '
    'retinal images. Nature, 622, 156\u2013163.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)


# ── Add page numbers ───────────────────────────────────────────────
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add PAGE field
    run = p.add_run()
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fld_char_begin)
    run2 = p.add_run()
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instr)
    run3 = p.add_run()
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fld_char_end)
    for r in [run, run2, run3]:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)


# ── Save ────────────────────────────────────────────────────────────
output_path = '/home/joneill/Nextcloud/vaults/jmind/calmi2/envision-discovery/paper/envision_discovery_paper.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
